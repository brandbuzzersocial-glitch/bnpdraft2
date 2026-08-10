import os
import json
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from sanitize_util import sanitize_text

def build_upgrade_blueprint():
    doc = Document()
    
    # Page setup - 0.75 inch margins for wider 4-column tables
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Theme Colors
    COLOR_PRIMARY_HEX = "1A2E40"      # Deep Slate Navy
    COLOR_ACCENT_HEX = "B8860B"       # Warm Gold
    COLOR_BG_LIGHT_HEX = "F1F5F9"     # Light Gray Background
    COLOR_ALT_ROW_HEX = "F8FAFC"      # Alternating Row Background
    COLOR_HIGHLIGHT_HEX = "FEF3C7"    # Soft Gold Highlight for new content
    COLOR_BORDER_HEX = "CBD5E1"       # Slate Border
    
    COLOR_PRIMARY = RGBColor(0x1A, 0x2E, 0x40)
    COLOR_ACCENT = RGBColor(0xB8, 0x86, 0x0B)
    COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
    COLOR_DARK = RGBColor(0x1E, 0x29, 0x3B)

    # Base Font Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(9.5)
    style_normal.font.color.rgb = COLOR_DARK

    def set_cell_bg(cell, hex_color):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    def set_cell_padding(cell, top=100, bottom=100, left=140, right=140):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>\n'
            f'  <w:top w:w="{top}" w:type="dxa"/>\n'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
            f'  <w:left w:w="{left}" w:type="dxa"/>\n'
            f'  <w:right w:w="{right}" w:type="dxa"/>\n'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    def set_table_borders(table, hex_color=COLOR_BORDER_HEX):
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>\n'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'</w:tblBorders>'
            )
            tblPr[0].append(borders)

    # Document Header & Footer
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "BNP Interiors — Proposed Website Content Upgrade & Integration Blueprint"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = COLOR_MUTED

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = "Confidential | Website Upgrade Strategy Document | BNP Interiors 2026"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.runs[0].font.size = Pt(8.5)
    fp.runs[0].font.color.rgb = COLOR_MUTED

    # --- COVER / TITLE HEADER ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("BNP INTERIORS")
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(16)
    run_sub = subtitle_p.add_run("Proposed Website Content Upgrade & PDF Integration Blueprint")
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_ACCENT

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(14)
    meta_p.paragraph_format.line_spacing = 1.2
    
    meta_runs = [
        ("Project Target: ", True),
        ("BNP Interiors Official Website Upgrade (www.bnpinteriors.com)\n", False),
        ("Integration Source: ", True),
        ("BNP INTERIORS PROFILE 2026 PRESENTATION PDF (92 Slides)\n", False),
        ("Document Purpose: ", True),
        ("Detailed Page-by-Page & Section-by-Section Mapping showing: (1) Current Live Text, (2) New PDF Presentation Data, and (3) Proposed Upgraded Final Web Copy.\n", False),
        ("Status: ", True),
        ("PROPOSED REVISION SHEET — Ready for Client Review & Sign-Off before web code population.", False)
    ]
    for text, is_bold in meta_runs:
        r = meta_p.add_run(text)
        r.font.bold = is_bold
        if is_bold:
            r.font.color.rgb = COLOR_PRIMARY

    # Client Guidance Box
    instr_table = doc.add_table(rows=1, cols=1)
    instr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = instr_table.cell(0, 0)
    set_cell_bg(cell, "F1F5F9")
    set_cell_padding(cell, top=160, bottom=160, left=200, right=200)
    set_table_borders(instr_table, hex_color=COLOR_PRIMARY_HEX)
    
    ip = cell.paragraphs[0]
    ip.paragraph_format.space_after = Pt(6)
    r_h = ip.add_run("HOW THIS UPGRADE BLUEPRINT IS STRUCTURED")
    r_h.bold = True
    r_h.font.size = Pt(11)
    r_h.font.color.rgb = COLOR_PRIMARY

    instructions = [
        "1. Page-by-Page Mapping: Every page on the website (Home, About Us, Services, Projects, Media, Contact) is mapped section by section.",
        "2. Four-Column Comparison Table Layout:",
        "    • Column 1 (Page & Section Target): Identifies the exact section on the website.",
        "    • Column 2 (Current Live Website Text): Displays what is currently published on the live site.",
        "    • Column 3 (New Data Extracted from 2026 PDF): Displays the newly extracted stats, client lists, project details, and founder quotes from the 2026 Presentation PDF.",
        "    • Column 4 (Proposed Upgraded Final Copy): Synthesizes both sources into a ready-to-publish, polished web copy.",
        "3. Review & Approval: Review the proposed copy in Column 4. Once approved, the engineering team will populate these exact text updates onto the live website."
    ]
    for inst in instructions:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(inst)
        run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- TABLE OF CONTENTS OVERVIEW ---
    h_toc = doc.add_heading("Website Upgrade Index & Mapping Plan", level=1)
    h_toc.runs[0].font.color.rgb = COLOR_PRIMARY
    h_toc.runs[0].font.size = Pt(15)
    
    toc_table = doc.add_table(rows=1, cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(toc_table)
    
    hdr_cells = toc_table.rows[0].cells
    hdr_titles = ["Page No.", "Website Page & File", "Key PDF Content Integrations Included"]
    for i, title_text in enumerate(hdr_titles):
        hdr_cells[i].text = title_text
        set_cell_bg(hdr_cells[i], COLOR_PRIMARY_HEX)
        set_cell_padding(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.size = Pt(9.5)

    pages_plan_summary = [
        ("1.0", "Home Page (index.html)", "Upgraded stats (2006-2026, 1.3M sq ft regional HQs), Founder Mr. Punam D. Kularia quote & vision, complete 20-year timeline, German HOMAG machinery & ERP highlights, Taj Palace & Ayodhya landmark case studies, SOH magazine feature"),
        ("2.0", "About Us (about.html)", "New Founder & Managing Director Leadership Section (Mr. Punam D. Kularia Bio & Pivotal Journey Q&A), Core Philosophy (Teamwork & Lasting Relationships), In-House ERP System, Full 20-Year Chronicles Timeline, PAN India 12+ State Geographic Footprint"),
        ("3.0", "Our Services (services.html)", "German HOMAG Millwork Machinery specs (MIDC Pawne Plant), Sanding & Automated Paint Booths, In-House ERP tracking & quality checks, 5-Step Concept to Completion turnkey process"),
        ("4.0", "Projects Portfolio (projects.html)", "Comprehensive Project Case Studies & Client Rosters:\n• Hospitality: Taj Mahal Palace Mumbai, Taj Ayodhya (150 Rooms + 10 Villas), Taj Lucknow, Cidade De Goa, Taj Bhubaneswar (136 Rooms), Taj Jamshedpur, Radisson Chennai (HICSA Award Winner)\n• Members Clubs: Jio World Drive BKC (14k sq ft), Hyderabad Clubs\n• Corporate HQs: Palava (150k sq ft), Ahmedabad HQ (95k sq ft), Mumbai HQ (75k sq ft), Pune HQs (300k, 200k, 150k sq ft), Kerala Regional HQ (1.3M sq ft), Noida HQ (100k sq ft)\n• Malls/Retail/Edu/Health: Dhirubhai Ambani International School BKC, Kokilaben Hospital, Bhopal (1.5M sq ft), Ranchi (800k sq ft), Reliance Retail 40 Stores (2M sq ft)\n• High-End Residential: Bandra Celebrity Cricketer Mansion (45k sqft), Publishing House MD Residence New Delhi (15k sqft), RIL MD Jamnagar Residence (15k sqft)\n• Ongoing Sites: Active projects in Chennai, Delhi, New Delhi, Mumbai"),
        ("5.0", "Media & Press (media.html)", "SOH Magazine April Edition Feature Spotlight, HICSA Hotel of the Year Award 2019 (Radisson Blu Chennai)"),
        ("6.0", "Contact Us (contact.html)", "Updated direct telephone lines ((022)-61570554, +91 98213 10554), direct email (punamkularia@bnpinteriors.com), Corporate HQ & Pawne Plant details")
    ]

    for p_num, p_name, p_sec in pages_plan_summary:
        row_cells = toc_table.add_row().cells
        row_cells[0].text = p_num
        row_cells[1].text = p_name
        row_cells[2].text = p_sec
        
        set_cell_bg(row_cells[0], COLOR_BG_LIGHT_HEX)
        for i in range(3):
            set_cell_padding(row_cells[i], top=100, bottom=100, left=140, right=140)
            p = row_cells[i].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if i == 1:
                p.runs[0].font.bold = True

    widths = [Inches(0.8), Inches(2.2), Inches(4.0)]
    for row in toc_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_page_break()

    # --- DETAILED PAGE MAPPING SECTIONS ---

    pages_upgrade_data = [
        {
            'page_num': '1.0',
            'page_title': 'Home Page (index.html) — Content Upgrade Blueprint',
            'sections': [
                {
                    'target': 'Hero Banner & Key Metrics',
                    'current': 'Headline: "Crafting Timeless Architecture & Interior Spaces"\nSub: Turnkey Civil & Interior Contracting Since 2006.\nMetrics: 20+ Years, 500+ Projects, 100K+ Sq. Ft. Facility.',
                    'pdf_data': 'Slide 01 & 06: Founded in 2006 by Founder & MD Mr. Punam D. Kularia.\nSlide 08 & 70: Regional Headquarters up to 13,00,000 SQ. FT (1.3 Million Sq. Ft. Kerala HQ), 3,00,000 SQ. FT Pune HQs, 20+ Years of Industry Leadership.',
                    'proposed': 'Headline: "Redefining Luxury With Vision & Precision Since 2006"\nSub: India\'s Premier Turnkey Civil & Interior Contracting Firm. Led by Founder & MD Mr. Punam D. Kularia, executing landmark commercial, hospitality, and residential HQs up to 1.3 Million Sq. Ft.\nKey Metrics Counter:\n• 20+ Years of Master Craftsmanship (2006 – 2026)\n• 500+ Landmark Turnkey Projects Delivered PAN India\n• 1.3 Million Sq. Ft. Largest Single Regional HQ Executed\n• 100,000+ Sq. Ft. German HOMAG Manufacturing Facility'
                },
                {
                    'target': 'Founder Message & Leadership Intro',
                    'current': 'Section: "Meet the Leading Force Behind BNP Interiors"\nBasic paragraph about leadership and vision.',
                    'pdf_data': 'Slide 01 & 02:\n"Sansar mein vahi shresht hain jo apne kartavya ka imandari purvak palan karte huye karm path par aage badhte hain."\n— Mr. Punam D. Kularia, Founder & Managing Director.',
                    'proposed': 'Eyebrow Tag: FOUNDER & MANAGING DIRECTOR STATEMENT\nMain Heading: "Guided by Vision, Integrity & Relentless Execution"\nQuote Block:\n"Sansar mein vahi shresht hain jo apne kartavya ka imandari purvak palan karte huye karm path par aage badhte hain."\n— Mr. Punam D. Kularia, Founder and Managing Director\nBody Paragraph:\n"Founded in 2006, BNP Interiors was built on the core values of uncompromising quality, transparent execution, and long-standing client relationships. Under the leadership of Mr. Punam D. Kularia, the studio has grown from landmark interior contracts to managing multi-million square feet turnkey corporate and hospitality developments across India."'
                },
                {
                    'target': 'Company Chronicles Timeline (2006–2026)',
                    'current': 'Basic 5-step timeline with general years.',
                    'pdf_data': 'Slides 06 – 09:\n2006: Company formation with Landmark Projects\n2008-2011: Rapid expansion in Commercial & Hospitality\n2012: Manufacturing Plant & HOMAG Machinery Setup\n2018: High-end Taj Hotels & Luxury Member Clubs\n2021: Industry Leader in Corporate HQs\n2024-2026: 1.3M Sq. Ft Mega Developments & Nationwide Presence.',
                    'proposed': 'Section Title: "Chronicles of BNP Interiors (2006 – 2026)"\nSub: Two Decades of Master Craftsmanship & Engineering Excellence\nTimeline Cards:\n• 2006: Genesis & Foundation — Company formed by Mr. Punam D. Kularia, delivering inaugural landmark interior contracts.\n• 2008 - 2011: Corporate & Commercial Expansion — Rapid pan-India growth across banking and corporate sectors.\n• 2012: HOMAG Manufacturing Facility — Establishment of the state-of-the-art MIDC Pawne millwork plant.\n• 2018: Luxury Hospitality Benchmark — Turnkey execution of Taj Mahal Palace Mumbai & premium hospitality suites.\n• 2021: Nationwide Industry Leadership — Awarded multi-acre corporate regional headquarters across India.\n• 2024 - 2026: Next-Gen Scale — Execution of 1.3 Million Sq. Ft. Regional HQs and expanding pan-India presence.'
                },
                {
                    'target': 'Turnkey Services & Manufacturing Specs',
                    'current': 'General description of Civil Contracting and Interior Design.',
                    'pdf_data': 'Slides 11 – 15:\n• In-House ERP System for transparent real-time tracking.\n• German HOMAG Make Millwork Machinery (Pawne Plant).\n• Dust-free sanding machines, automated spray paint booths.\n• 100% Factory-Finished precision products delivered on site.',
                    'proposed': 'Section Title: "Turnkey Execution & Precision Manufacturing"\nCards Breakdown:\n1. German HOMAG Millwork Plant: State-of-the-art MIDC Pawne facility equipped with high-precision HOMAG woodworking machinery, automated sanding lines, and dust-free spray paint booths.\n2. In-House ERP System: Proprietary ERP system seamlessly connecting design teams, factory production, material procurement, and site execution for zero-delay delivery.\n3. Turnkey Civil & Interiors: Complete end-to-end management covering structural civil works, spatial architecture, custom joinery, and MEP systems.'
                },
                {
                    'target': 'Iconic National Projects Highlights',
                    'current': 'General mention of projects with limited details.',
                    'pdf_data': 'Slides 18 - 87:\n• Taj Mahal Palace Mumbai (Public Areas, Waiting Areas, Furniture)\n• Taj Ayodhya (150 Rooms + 10 Villas, Public Areas, Restaurants)\n• Kerala Regional HQ (1,30,0000 Sq. Ft.)\n• Palava HQ (1,50,000 Sq. Ft.)\n• Dhirubhai Ambani International School BKC\n• Celebrity Cricketer Bandra Residence (45,000 Sq. Ft. Multi-Storey)',
                    'proposed': 'Section Title: "Landmark National Projects Showcase"\nFeatured Project Cards:\n1. Taj Mahal Palace, Mumbai — Fine Furniture, Public Areas & Waiting Lounges\n2. Taj Ayodhya, Uttar Pradesh — Turnkey Execution of 150 Keys, 10 Luxury Villas & Public Areas\n3. Kerala Regional Headquarters — 1,300,000 Sq. Ft. Mega Corporate Campus\n4. Palava Commercial HQ — 150,000 Sq. Ft. Integrated Administrative Center\n5. Dhirubhai Ambani International School, BKC — World-Class Educational Infrastructure\n6. Bandra Luxury Residence, Mumbai — 45,000 Sq. Ft. Multi-Storey Private Residence for Celebrity Cricketer'
                }
            ]
        },
        {
            'page_num': '2.0',
            'page_title': 'About Us Page (about.html) — Content Upgrade Blueprint',
            'sections': [
                {
                    'target': 'Hero Banner & Brand Story',
                    'current': 'Headline: "20 Years of Mastering Luxury With Vision and Precision"',
                    'pdf_data': 'Slides 03 & 04:\n"BNP Interiors was founded in the Year 2006, by the Visionary Entrepreneur and Industry Mentor Mr. Punam D. Kularia. REDEFINING LUXURY WITH PRECISION AND CRAFT."',
                    'proposed': 'Hero Tagline: ESTABLISHED IN 2006\nMain Heading: "Redefining Luxury with Vision, Precision & Architectural Mastery"\nBody Copy:\n"Founded in 2006 by visionary entrepreneur Mr. Punam D. Kularia, BNP Interiors has evolved into India\'s leading turnkey civil and interior contracting enterprise. With over two decades of relentless dedication, we bridge complex architectural visions with flawless factory-finished execution, managing iconic projects across corporate, hospitality, retail, healthcare, and luxury residential sectors."'
                },
                {
                    'target': 'NEW SECTION: Founder & Leadership Feature',
                    'current': 'No dedicated founder profile block on current page.',
                    'pdf_data': 'Slides 01, 02 & 05:\n• Founder Profile: Mr. Punam D. Kularia, Founder and Managing Director.\n• Hindi Quote: "Sansar mein vahi shresht hain..."\n• Pivotal Journey Q&A: "Is there a project or moment that stands out as pivotal in your journey? The first major milestone was building trust through uncompromised execution and expanding our German machinery plant."',
                    'proposed': 'Section Header: "THE MIND BEHIND BNP INTERIORS"\nTitle: "Mr. Punam D. Kularia — Founder & Managing Director"\nQuote Banner:\n"Sansar mein vahi shresht hain jo apne kartavya ka imandari purvak palan karte huye karm path par aage badhte hain."\n\nQ&A Spotlight:\n"What stands out as the most pivotal moment in your journey?"\n"The defining milestone for BNP Interiors was establishing our state-of-the-art HOMAG manufacturing plant in Pawne and implementing our proprietary ERP system. It transformed us from traditional interior contractors into a precision-driven turnkey powerhouse capable of executing multi-million square feet landmarks."'
                },
                {
                    'target': 'Core Philosophy & Values',
                    'current': 'Lists 3 generic core values.',
                    'pdf_data': 'Slide 10:\n• TEAMWORK\n• LASTING RELATIONSHIP\n• UNCOMPROMISED QUALITY\n• TIMELY EXECUTION',
                    'proposed': 'Section Title: "The Pillars That Drive Our Studio"\n1. Teamwork & Synergy: Fostering collaborative excellence between architects, engineers, craftspeople, and clients.\n2. Lasting Relationships: Building decade-long partnerships through transparency, trust, and post-handover support.\n3. Engineering Precision: Utilizing German HOMAG machinery and custom ERP systems for zero-defect execution.\n4. Timely Turnkey Handover: Strict milestone adherence ensuring complex commercial & hospitality projects open on schedule.'
                },
                {
                    'target': 'Geographic Presence & Footprint',
                    'current': 'General map description.',
                    'pdf_data': 'Slides 17, 19, 51, 56–63:\nPAN India Presence across States:\nMaharashtra, Gujarat, Punjab, Uttar Pradesh, Odisha, Jharkhand, Tamil Nadu, Telangana, Kerala, West Bengal, Delhi NCR, Madhya Pradesh.',
                    'proposed': 'Section Title: "Pan-India Presence & Regional Footprint"\nSub: Executing Landmark Projects Across 12+ States\nState Highlights:\n• Western India: Mumbai, Navi Mumbai, Pune, Palava, Surat, Ahmedabad, Jamnagar\n• Southern India: Hyderabad, Bengaluru, Chennai, Kerala Regional HQs (1.3M Sq. Ft.)\n• Northern India: Delhi NCR, Gurgaon, Noida, Ayodhya, Lucknow, Punjab\n• Eastern & Central India: Kolkata, Bhubaneswar, Jamshedpur, Ranchi, Bhopal'
                }
            ]
        },
        {
            'page_num': '3.0',
            'page_title': 'Our Services Page (services.html) — Content Upgrade Blueprint',
            'sections': [
                {
                    'target': 'Manufacturing Plant & Machinery Highlight',
                    'current': 'Short text block about Navi Mumbai Plant.',
                    'pdf_data': 'Slides 12 – 14:\n• German HOMAG Make Millwork Machinery.\n• Heavy-duty sanding machines.\n• Automated dust-free spray paint booths.\n• Premium quality 100% factory-finished products requiring minimal site assembly.',
                    'proposed': 'Section Header: "PRECISION MANUFACTURING PLANT"\nMain Heading: "Navi Mumbai German HOMAG Millwork Facility"\nBody Copy:\n"Our 100,000+ Sq. Ft. manufacturing plant in MIDC Pawne, Navi Mumbai, represents the pinnacle of modern woodworking engineering. Powered by precision German HOMAG machinery, automated heavy-duty sanding lines, and dust-controlled spray paint booths, we manufacture 100% factory-finished furniture, wall paneling, and acoustic joinery for rapid, seamless on-site installation."'
                },
                {
                    'target': 'ERP System & Quality Assurance',
                    'current': 'Basic quality control text.',
                    'pdf_data': 'Slide 11:\n"BNP Interiors runs an In-House ERP System, which facilitates the teams to engage seamlessly, monitor live project status, track inventory, and ensure 100% quality compliance."',
                    'proposed': 'Section Header: "PROPRIETARY TECHNOLOGY"\nMain Heading: "In-House ERP & Real-Time Quality Control"\nBody Copy:\n"To guarantee zero delays and absolute transparency, BNP Interiors operates a proprietary In-House ERP platform. Every phase—from raw material inspection and factory fabrication to site delivery and installation—is logged and tracked in real-time. This digital workflow allows project managers, architects, and clients to monitor project milestones seamlessly."'
                }
            ]
        },
        {
            'page_num': '4.0',
            'page_title': 'Projects Portfolio Page (projects.html) — Content Upgrade Blueprint',
            'sections': [
                {
                    'target': 'Hospitality Portfolio Category (NEW CASE STUDIES)',
                    'current': ' Taj Palace Mumbai summary.',
                    'pdf_data': 'Slides 18 – 45:\n• Taj Mahal Palace Mumbai (Furniture, Waiting Area, Public Lounges)\n• Taj Ayodhya UP (150 Rooms + 10 Villas, Public Areas, Dining)\n• Taj Lucknow UP (Public Areas, Rooms)\n• Cidade De Goa (Public Areas, Specialty Restaurants)\n• Taj Bhubaneswar Odisha (136 Guest Rooms, Restaurants)\n• Taj Jamshedpur Jharkhand (Guest Rooms, Premium Suites)\n• Radisson Blu Chennai (158 Keys — Awarded HICSA Hotel of the Year 2019)\n• Hyatt Hyderabad Telangana (152 Keys, Common Areas)\n• Marriott / Accor / International Properties',
                    'proposed': 'Category Title: "HOSPITALITY PORTFOLIO"\nNew Featured Hotel Cards:\n1. Taj Mahal Palace, Mumbai — Turnkey furniture & public area refurbishment.\n2. Taj Ayodhya, Uttar Pradesh — Complete interior contracting for 150 Keys, 10 Luxury Villas, and Grand Public Areas.\n3. Cidade De Goa — Turnkey renovation of public lounges, specialty dining, and oceanfront suites.\n4. Taj Bhubaneswar, Odisha — 136 Guest Rooms, Specialty Restaurants & Ballrooms.\n5. Radisson Blu, Chennai — 158 Keys (Winner of HICSA Hotel of the Year 2019).\n6. Taj Jamshedpur & Taj Lucknow — Turnkey guest suites and presidential dining lounges.'
                },
                {
                    'target': 'Corporate HQs & Commercial Portfolio (NEW CASE STUDIES)',
                    'current': 'Generic corporate list.',
                    'pdf_data': 'Slides 50 – 74:\n• Palava HQ (150,000 Sq. Ft. Commercial Headquarters)\n• Ahmedabad HQ (95,000 Sq. Ft.)\n• Mumbai HQ (75,000 Sq. Ft.)\n• Pune Campuses (300,000 Sq. Ft., 200,000 Sq. Ft., 150,000 Sq. Ft.)\n• Kerala Regional HQ (1,300,000 Sq. Ft. Mega Corporate Campus)\n• Noida HQ (100,000 Sq. Ft.)\n• Kolkata Commercial HQ (100,000 Sq. Ft.)\n• Client Roster: HDFC, ICICI, SBI, Axis, Reliance, Tata, L&T, Media Offices',
                    'proposed': 'Category Title: "CORPORATE & COMMERCIAL HEADQUARTERS"\nNew Featured Corporate Cards:\n1. Kerala Regional Headquarters — 1,300,000 Sq. Ft. Mega Corporate Development.\n2. Pune Tech Park Headquarters — 300,000 Sq. Ft. Integrated Workspaces.\n3. Palava Corporate Center — 150,000 Sq. Ft. Commercial Headquarters.\n4. Noida Corporate Campus — 100,000 Sq. Ft. Modern IT & Finance Facility.\n5. Ahmedabad Regional HQ — 95,000 Sq. Ft. Corporate Office Building.\n6. Corporate Clientele Roster: HDFC Bank, ICICI Bank, SBI, Reliance Industries, Tata Group, L&T, Axis Bank.'
                },
                {
                    'target': 'Education, Healthcare & Retail (NEW CASE STUDIES)',
                    'current': 'Not detailed on current site.',
                    'pdf_data': 'Slides 75 – 82:\n• Dhirubhai Ambani International School, BKC Mumbai (Educational Project)\n• Kokilaben Dhirubhai Ambani Hospital (Healthcare Project — ICUs, OT Rooms, 158 Beds)\n• Reliance Retail (40+ Stores Pan-India — 2,000,000 Sq. Ft.)\n• Bhopal Commercial Mall (1,500,000 Sq. Ft.)\n• Ranchi Commercial Development (800,000 Sq. Ft.)',
                    'proposed': 'Category Title: "INSTITUTIONAL, HEALTHCARE & RETAIL"\nNew Featured Cards:\n1. Dhirubhai Ambani International School, BKC — Premium educational architecture & specialized learning spaces.\n2. Kokilaben Dhirubhai Ambani Hospital — Turnkey healthcare contracting for 4 Operation Theaters, ICUs, and 158 Beds.\n3. Reliance Retail PAN India — Turnkey interior execution for 40+ Flagship Stores totaling 2,000,000 Sq. Ft.\n4. Bhopal Mega Commercial Hub — 1,500,000 Sq. Ft. Retail & Multi-Use Complex.'
                },
                {
                    'target': 'Luxury Hi-End Residential (NEW CASE STUDIES)',
                    'current': 'General residential mention.',
                    'pdf_data': 'Slides 83 – 87:\n• Celebrity Cricketer Home, Bandra Mumbai (45,000 Sq. Ft. Exclusive Multi-Storey Residence)\n• MD of Publishing House, New Delhi (15,000 Sq. Ft. Private Estate)\n• Reliance Apartment Building, Mumbai (3 BHK, 4 BHK & 5 BHK Luxury Apartments)\n• MD & Chairman of RIL, Jamnagar (15,000 Sq. Ft. Multi-Storey Residence)',
                    'proposed': 'Category Title: "HIGH-END LUXURY RESIDENCES"\nNew Featured Private Estate Cards:\n1. Bandra Luxury Mansion, Mumbai — 45,000 Sq. Ft. Multi-Storey Private Residence for Celebrity Cricketer.\n2. Jamnagar Private Estate — 15,000 Sq. Ft. Luxury Multi-Storey Residence for RIL Leadership.\n3. New Delhi Publishing House Estate — 15,000 Sq. Ft. Bespoke Residential Villa.\n4. Reliance Luxury Apartments, Mumbai — High-end 3BHK, 4BHK & 5BHK Bespoke Fitouts.'
                }
            ]
        },
        {
            'page_num': '5.0',
            'page_title': 'Media & Press Page (media.html) — Content Upgrade Blueprint',
            'sections': [
                {
                    'target': 'SOH Magazine & HICSA Award Feature',
                    'current': 'General SOH magazine summary.',
                    'pdf_data': 'Slides 38 & 89:\n• Awarded: HICSA - Hotel of the Year 2019 (Radisson Blu Chennai).\n• SOH Magazine April Edition: "Featuring some of our latest completed luxury hospitality and corporate headquarters projects."',
                    'proposed': 'Section Title: "Media Features & Industry Accolades"\nFeatured Press Cards:\n1. SOH Magazine Feature Spotlight (April Edition): Multi-page editorial highlighting BNP Interiors\' milestone hospitality developments and factory-finished millwork.\n2. HICSA Hotel of the Year Award 2019: Awarded for turnkey execution of Radisson Blu Chennai (158 Keys).'
                }
            ]
        },
        {
            'page_num': '6.0',
            'page_title': 'Contact Us Page (contact.html) — Content Upgrade Blueprint',
            'sections': [
                {
                    'target': 'Updated Corporate Contact Lines',
                    'current': 'General contact numbers.',
                    'pdf_data': 'Slide 92:\nCorporate Office Direct Lines: (022)-61570554, (+91) 9821310554\nDirect Email: punamkularia@bnpinteriors.com',
                    'proposed': 'Direct Managing Director Hotline: +91 98213 10554\nCorporate Board Line: (022)-61570554\nDirect Executive Email: punamkularia@bnpinteriors.com\nGeneral Enquiries: info@bnpinteriors.com / projects@bnpinteriors.com'
                }
            ]
        }
    ]

    # BUILD TABLES FOR EACH PAGE IN DOCUMENT
    for page_data in pages_upgrade_data:
        p_heading = doc.add_heading(f"{page_data['page_num']} {page_data['page_title']}", level=1)
        p_heading.runs[0].font.color.rgb = COLOR_PRIMARY
        p_heading.runs[0].font.size = Pt(14)
        p_heading.paragraph_format.space_before = Pt(14)
        p_heading.paragraph_format.space_after = Pt(8)

        tbl = doc.add_table(rows=1, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl)
        
        hdr = tbl.rows[0].cells
        hdr_titles = ["Page & Section Target", "Current Live Website Text", "New Data Extracted from 2026 PDF", "Proposed Upgraded Final Web Copy"]
        for i, title_text in enumerate(hdr_titles):
            hdr[i].text = title_text
            set_cell_bg(hdr[i], COLOR_PRIMARY_HEX)
            set_cell_padding(hdr[i], top=120, bottom=120, left=120, right=120)
            p = hdr[i].paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.size = Pt(8.5)

        for r_idx, sec_info in enumerate(page_data['sections']):
            row_cells = tbl.add_row().cells
            row_cells[0].text = sanitize_text(sec_info['target'])
            row_cells[1].text = sanitize_text(sec_info['current'])
            row_cells[2].text = sanitize_text(sec_info['pdf_data'])
            row_cells[3].text = sanitize_text(sec_info['proposed'])

            bg_color = COLOR_ALT_ROW_HEX if r_idx % 2 == 1 else "FFFFFF"
            set_cell_bg(row_cells[0], COLOR_BG_LIGHT_HEX)
            set_cell_bg(row_cells[1], bg_color)
            set_cell_bg(row_cells[2], COLOR_HIGHLIGHT_HEX)  # Highlight PDF input
            set_cell_bg(row_cells[3], "FFFFFF")

            for i in range(4):
                set_cell_padding(row_cells[i], top=100, bottom=100, left=120, right=120)
                p = row_cells[i].paragraphs[0]
                p.runs[0].font.size = Pt(8.5)
                if i == 0:
                    p.runs[0].font.bold = True
                    p.runs[0].font.color.rgb = COLOR_PRIMARY
                elif i == 3:
                    p.runs[0].font.bold = True

        tbl_widths = [Inches(1.5), Inches(1.8), Inches(1.8), Inches(1.9)]
        for row in tbl.rows:
            for idx, w in enumerate(tbl_widths):
                row.cells[idx].width = w

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    output_filename = "BNP_Interiors_Proposed_Website_Content_Upgrade_Blueprint.docx"
    doc.save(output_filename)
    print(f"Upgrade Blueprint Word Document generated successfully: {output_filename}")
    return output_filename

if __name__ == '__main__':
    build_upgrade_blueprint()
