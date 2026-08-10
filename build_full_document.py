import os
import glob
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from extract_all_content import extract_page_elements
from sanitize_util import sanitize_text

def build_word_document():
    doc = Document()
    
    # 1 Inch Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Theme Colors
    COLOR_PRIMARY_HEX = "1A2E40"      # Deep Navy Blue
    COLOR_ACCENT_HEX = "B8860B"       # Warm Gold / Dark Goldenrod
    COLOR_BG_LIGHT_HEX = "F1F5F9"     # Soft Gray background
    COLOR_ALT_ROW_HEX = "F8FAFC"      # Subtle alternating row background
    COLOR_BORDER_HEX = "CBD5E1"       # Border slate gray
    
    COLOR_PRIMARY = RGBColor(0x1A, 0x2E, 0x40)
    COLOR_ACCENT = RGBColor(0xB8, 0x86, 0x0B)
    COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
    COLOR_DARK = RGBColor(0x1E, 0x29, 0x3B)

    # Base Font Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = COLOR_DARK

    def set_cell_bg(cell, hex_color):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    def set_cell_padding(cell, top=120, bottom=120, left=160, right=160):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>\n'
            f'  <w:top w:w="{top}" w:type="dxa"/>\n'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
            f'  <w:left w:w="{left}" w:type="dxa"/>\n'
            f'  <w:right w:w="{right}" w:type="dxa"/>\n'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    def set_table_borders(table, hex_color=COLOR_BORDER_HEX):
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>\n'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'</w:tblBorders>'
            )
            tblPr[0].append(borders)

    # Document Header & Footer
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "BNP Interiors — Complete Website Copywriting & Client Revision Document"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = COLOR_MUTED

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = "Confidential | Website Copy Review Document | BNP Interiors"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.runs[0].font.size = Pt(8.5)
    fp.runs[0].font.color.rgb = COLOR_MUTED

    # --- COVER / TITLE HEADER ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("BNP INTERIORS")
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(16)
    run_sub = subtitle_p.add_run("Complete Website Content Copywriting & Client Revision Document")
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_ACCENT

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(14)
    meta_p.paragraph_format.line_spacing = 1.2
    
    meta_runs = [
        ("Project Name: ", True),
        ("BNP Interiors Official Website (www.bnpinteriors.com)\n", False),
        ("Document Scope: ", True),
        ("Complete Section-by-Section & Page-by-Page Text Content Breakdown\n", False),
        ("Target Audience: ", True),
        ("Client Review Team, Content Copywriter, & Project Management\n", False),
        ("Purpose: ", True),
        ("To facilitate easy client edits, content revisions, copy approvals, and section updates.", False)
    ]
    for text, is_bold in meta_runs:
        r = meta_p.add_run(text)
        r.font.bold = is_bold
        if is_bold:
            r.font.color.rgb = COLOR_PRIMARY

    # Client Guidance / Instructions Box
    instr_table = doc.add_table(rows=1, cols=1)
    instr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = instr_table.cell(0, 0)
    set_cell_bg(cell, "F1F5F9")
    set_cell_padding(cell, top=160, bottom=160, left=200, right=200)
    set_table_borders(instr_table, hex_color=COLOR_PRIMARY_HEX)
    
    ip = cell.paragraphs[0]
    ip.paragraph_format.space_after = Pt(6)
    r_h = ip.add_run("INSTRUCTIONS FOR CLIENT REVIEW & EDITING")
    r_h.bold = True
    r_h.font.size = Pt(11)
    r_h.font.color.rgb = COLOR_PRIMARY

    instructions = [
        "1. Structure: Every live page on the website is organized into clear visual sections (Hero Banner, Features, Services, Process, Case Studies, Testimonials, Contact Forms, etc.).",
        "2. Current Copy: The 'Current Website Text' column contains the exact live text currently displayed on the website.",
        "3. Providing Feedback & Changes:",
        "    - Option A (Recommended): Type your new text or modified copy directly into the 'Proposed New Text / Client Edits' column.",
        "    - Option B: Turn on Microsoft Word 'Track Changes' (Review tab > Track Changes) and edit the text directly.",
        "4. Unchanged Sections: Leave the 'Proposed New Text' cell blank for any text you wish to keep as-is.",
        "5. Next Steps: Save this document upon completion and return it to the project coordinator."
    ]
    for inst in instructions:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(inst)
        run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- TABLE OF CONTENTS OVERVIEW ---
    h_toc = doc.add_heading("Website Page Index & Structure Overview", level=1)
    h_toc.runs[0].font.color.rgb = COLOR_PRIMARY
    h_toc.runs[0].font.size = Pt(15)
    
    toc_table = doc.add_table(rows=1, cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(toc_table)
    
    hdr_cells = toc_table.rows[0].cells
    hdr_titles = ["Page No.", "Website Page Title & File", "Key Sections Included"]
    for i, title_text in enumerate(hdr_titles):
        hdr_cells[i].text = title_text
        set_cell_bg(hdr_cells[i], COLOR_PRIMARY_HEX)
        set_cell_padding(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.size = Pt(9.5)

    pages_summary = [
        ("0.0", "Global Navigation & Header", "Navbar Links, Phone Number, Email, Working Hours, Top Bar, Consultation CTA Button"),
        ("1.0", "Home Page (index.html)", "Hero Banner, Metrics Counter, Manufacturing Plant Intro, Leadership Intro, Journey Timeline (2006-2026), Turnkey Services, Execution Process, Iconic Projects, Media Features, Testimonials, Leadership Bios, CTA Banner, Studio Insights, Newsletter Signup"),
        ("2.0", "About Us (about.html)", "Hero Banner, 20-Year Legacy Story, Core Values, Vision & Mission Framework, Executive Leadership Team & Bios, CTA Banner"),
        ("3.0", "Our Services (services.html)", "Hero Banner, Turnkey Service Verticals, Precision Civil & Millwork Capabilities, HOMAG Navi Mumbai Plant Details, Step-by-Step Execution Process, CTA Banner"),
        ("4.0", "Projects Portfolio (projects.html)", "Hero Banner, Category Filters, Showcase Projects (Taj Palace, Luxury Commercial & Residential Projects), CTA Banner"),
        ("5.0", "Media & Press (media.html)", "Hero Banner, SOH Magazine Feature Spotlight, Editorial Summary, Media Coverage, Awards & Recognition"),
        ("6.0", "Blog & Insights (blog.html)", "Hero Banner, Featured Editorial on Sustainability & Biophilic Design (2027 Outlook), Architectural Trends"),
        ("7.0", "Contact Us (contact.html)", "Hero Banner, Corporate Office Details, Navi Mumbai Plant Location, Consultation Request Form Labels, Frequently Asked Questions (FAQs)"),
        ("8.0", "Global Footer & Legal Information", "Brand Tagline, Quick Navigation Links, Office Addresses, Direct Hotlines, Copyright & Disclaimer")
    ]

    for p_num, p_name, p_sec in pages_summary:
        row_cells = toc_table.add_row().cells
        row_cells[0].text = p_num
        row_cells[1].text = p_name
        row_cells[2].text = p_sec
        
        set_cell_bg(row_cells[0], COLOR_BG_LIGHT_HEX)
        for i in range(3):
            set_cell_padding(row_cells[i], top=100, bottom=100, left=140, right=140)
            p = row_cells[i].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if i == 1:
                p.runs[0].font.bold = True

    widths = [Inches(0.8), Inches(2.4), Inches(3.6)]
    for row in toc_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_page_break()

    # --- SECTION BUILDER HELPER ---
    def add_page_section_table(doc, page_num_str, page_title, page_sections):
        p_heading = doc.add_heading(f"{page_num_str} {page_title}", level=1)
        p_heading.runs[0].font.color.rgb = COLOR_PRIMARY
        p_heading.runs[0].font.size = Pt(15)
        p_heading.paragraph_format.space_before = Pt(16)
        p_heading.paragraph_format.space_after = Pt(10)

        for sec in page_sections:
            sec_name = sec['section_title']
            elements = sec['elements']
            
            if not elements:
                continue

            sh = doc.add_heading(f"Section: {sec_name}", level=2)
            sh.runs[0].font.color.rgb = COLOR_ACCENT
            sh.runs[0].font.size = Pt(12)
            sh.paragraph_format.space_before = Pt(12)
            sh.paragraph_format.space_after = Pt(6)

            tbl = doc.add_table(rows=1, cols=3)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(tbl)
            
            hdr = tbl.rows[0].cells
            hdr[0].text = "Element / Location"
            hdr[1].text = "Current Website Text"
            hdr[2].text = "Proposed New Text / Client Edits"
            
            for i in range(3):
                set_cell_bg(hdr[i], COLOR_PRIMARY_HEX)
                set_cell_padding(hdr[i], top=120, bottom=120, left=140, right=140)
                p = hdr[i].paragraphs[0]
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p.runs[0].font.size = Pt(9)

            for r_idx, (elem_type, current_txt) in enumerate(elements):
                row_cells = tbl.add_row().cells
                row_cells[0].text = sanitize_text(elem_type)
                row_cells[1].text = sanitize_text(current_txt)
                row_cells[2].text = ""  # Client edit area

                bg_color = COLOR_ALT_ROW_HEX if r_idx % 2 == 1 else "FFFFFF"
                set_cell_bg(row_cells[0], COLOR_BG_LIGHT_HEX)
                set_cell_bg(row_cells[1], bg_color)
                set_cell_bg(row_cells[2], "FFFFFF")

                for i in range(3):
                    set_cell_padding(row_cells[i], top=100, bottom=100, left=140, right=140)
                    p = row_cells[i].paragraphs[0]
                    p.runs[0].font.size = Pt(9)
                    if i == 0:
                        p.runs[0].font.bold = True
                        p.runs[0].font.color.rgb = COLOR_PRIMARY

            tbl_widths = [Inches(1.8), Inches(3.2), Inches(1.8)]
            for row in tbl.rows:
                for idx, w in enumerate(tbl_widths):
                    row.cells[idx].width = w
            
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 0. Global Navigation
    global_nav_sections = [{
        'section_title': 'Top Info Bar & Main Header Navigation',
        'elements': [
            ("Top Bar Tagline", "Premier Turnkey Civil & Interior Contracting Since 2006"),
            ("Phone Contact Numbers", "+91 98200 00000 / +91 22 2778 0000"),
            ("Email Address", "info@bnpinteriors.com / projects@bnpinteriors.com"),
            ("Operating Hours", "Mon - Sat: 9:30 AM - 6:30 PM IST"),
            ("Nav Link 1", "Home (index.html)"),
            ("Nav Link 2", "About Us (about.html)"),
            ("Nav Link 3", "Services (services.html)"),
            ("Nav Link 4", "Projects (projects.html)"),
            ("Nav Link 5", "Media & Press (media.html)"),
            ("Nav Link 6", "Blog (blog.html)"),
            ("Nav Link 7", "Contact Us (contact.html)"),
            ("Main Header CTA Button", "Request Consultation / Get In Touch")
        ]
    }]
    add_page_section_table(doc, "0.0", "Global Website Navigation & Header", global_nav_sections)
    doc.add_page_break()

    # HTML pages to process
    html_pages = [
        ("1.0", "Home Page", "index.html"),
        ("2.0", "About Us Page", "about.html"),
        ("3.0", "Our Services Page", "services.html"),
        ("4.0", "Projects Portfolio Page", "projects.html"),
        ("5.0", "Media & Press Page", "media.html"),
        ("6.0", "Blog & Insights Page", "blog.html"),
        ("7.0", "Contact Us Page", "contact.html")
    ]

    for page_num, page_title, file_name in html_pages:
        print(f"Parsing content from {file_name}...")
        sections_data = extract_page_elements(file_name)
        if sections_data:
            add_page_section_table(doc, page_num, f"{page_title} ({file_name})", sections_data)
            doc.add_page_break()

    # 8. Global Footer
    footer_sections = [{
        'section_title': 'Website Footer & Sub-Footer Content',
        'elements': [
            ("Footer Brand Heading", "BNP Interiors — Architecture & Turnkey Contracting"),
            ("Footer Brand Tagline", "Mastering Luxury with Vision and Precision. Delivering world-class interior contracting, spatial architecture, and custom millwork across India."),
            ("Corporate Address Title", "Corporate Head Office"),
            ("Corporate Office Address", "BNP Interiors, Prime Commercial Complex, Sector 15, CBD Belapur, Navi Mumbai, Maharashtra - 400614"),
            ("Manufacturing Facility Title", "Navi Mumbai Manufacturing Plant"),
            ("Manufacturing Plant Address", "Plot C-42, Industrial Area, MIDC Pawne, Navi Mumbai, Maharashtra - 400705"),
            ("Direct Phone Lines", "+91 98200 00000 | +91 22 2778 0000"),
            ("Official Email Enquiries", "info@bnpinteriors.com | projects@bnpinteriors.com"),
            ("Quick Links Column 1", "Home | About Us | Turnkey Services | Projects Portfolio"),
            ("Quick Links Column 2", "Media & SOH Magazine | Blog Insights | Contact Us | Privacy Policy"),
            ("Copyright Line", "© 2006 - 2026 BNP Interiors. All Rights Reserved. Crafted for Architectural Excellence.")
        ]
    }]
    add_page_section_table(doc, "8.0", "Global Website Footer & Legal Information", footer_sections)

    output_filename = "BNP_Interiors_Website_Content_Copy_Client_Review_Doc.docx"
    doc.save(output_filename)
    print(f"Word Document generated successfully: {output_filename}")
    return output_filename

if __name__ == '__main__':
    build_word_document()
