import os
import json
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from sanitize_util import sanitize_text

def build_presentation_word_document():
    with open('pdf_slides_extracted.json', 'r', encoding='utf-8') as f:
        slides = json.load(f)

    doc = Document()
    
    # 1 Inch Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Theme Colors
    COLOR_PRIMARY_HEX = "1A2E40"      # Deep Slate Navy
    COLOR_ACCENT_HEX = "B8860B"       # Warm Gold
    COLOR_BG_LIGHT_HEX = "F1F5F9"     # Light Gray Background
    COLOR_ALT_ROW_HEX = "F8FAFC"      # Alternating Row Background
    COLOR_BORDER_HEX = "CBD5E1"       # Slate Border
    
    COLOR_PRIMARY = RGBColor(0x1A, 0x2E, 0x40)
    COLOR_ACCENT = RGBColor(0xB8, 0x86, 0x0B)
    COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
    COLOR_DARK = RGBColor(0x1E, 0x29, 0x3B)

    # Base Font Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = COLOR_DARK

    def set_cell_bg(cell, hex_color):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    def set_cell_padding(cell, top=120, bottom=120, left=160, right=160):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>\n'
            f'  <w:top w:w="{top}" w:type="dxa"/>\n'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
            f'  <w:left w:w="{left}" w:type="dxa"/>\n'
            f'  <w:right w:w="{right}" w:type="dxa"/>\n'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    def set_table_borders(table, hex_color=COLOR_BORDER_HEX):
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>\n'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>\n'
                f'</w:tblBorders>'
            )
            tblPr[0].append(borders)

    # Document Header & Footer
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "BNP Interiors — Company Profile 2026 Presentation Content Extraction Sheet"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = COLOR_MUTED

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = "Confidential | 2026 Company Presentation PDF Copy Review | BNP Interiors"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.runs[0].font.size = Pt(8.5)
    fp.runs[0].font.color.rgb = COLOR_MUTED

    # --- COVER / TITLE HEADER ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("BNP INTERIORS")
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(16)
    run_sub = subtitle_p.add_run("Company Profile 2026 Presentation — Complete Content Extraction & Review Sheet")
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_ACCENT

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(14)
    meta_p.paragraph_format.line_spacing = 1.2
    
    meta_runs = [
        ("Source Document: ", True),
        ("BNP INTERIORS PROFILE 2026 PRESENTATION.pdf (92 Slides)\n", False),
        ("Document Scope: ", True),
        ("Complete Slide-by-Slide & Category Content Breakdown (Executive Profile, Timeline, Capabilities, Case Studies, Clients)\n", False),
        ("Target Audience: ", True),
        ("Client Review Team, Copywriting Team, & Web Strategy Team\n", False),
        ("Purpose: ", True),
        ("To review all rich content from the 2026 Presentation Deck and mark missing sections for website integration.", False)
    ]
    for text, is_bold in meta_runs:
        r = meta_p.add_run(text)
        r.font.bold = is_bold
        if is_bold:
            r.font.color.rgb = COLOR_PRIMARY

    # Client Guidance / Instructions Box
    instr_table = doc.add_table(rows=1, cols=1)
    instr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = instr_table.cell(0, 0)
    set_cell_bg(cell, "F1F5F9")
    set_cell_padding(cell, top=160, bottom=160, left=200, right=200)
    set_table_borders(instr_table, hex_color=COLOR_PRIMARY_HEX)
    
    ip = cell.paragraphs[0]
    ip.paragraph_format.space_after = Pt(6)
    r_h = ip.add_run("INSTRUCTIONS FOR CLIENT PRESENTATION COPY REVIEW")
    r_h.bold = True
    r_h.font.size = Pt(11)
    r_h.font.color.rgb = COLOR_PRIMARY

    instructions = [
        "1. Comprehensive Extraction: All text, client rosters, project highlights, milestones, equipment specs, and executive quotes have been extracted from all 92 slides of the 2026 Presentation Deck.",
        "2. Review Column Structure:",
        "    - Column 1 (Slide / Element Location): Specifies the Slide Number and Content Topic.",
        "    - Column 2 (Content Extracted from PDF Presentation): Contains the exact copy from the 2026 PDF.",
        "    - Column 3 (Proposed Website Placement / Client Notes): Dedicated blank column for your notes on where to place this content on the website (e.g., 'Add to About Page', 'Add as Case Study under Projects', 'Include in Services', or 'Do Not Include').",
        "3. Providing Edits: You can type your instructions directly into Column 3, or use Word Track Changes.",
        "4. Next Steps: Once reviewed, send this completed document back to proceed with populating the new content onto the website."
    ]
    for inst in instructions:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(inst)
        run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- TABLE OF CONTENTS OVERVIEW ---
    h_toc = doc.add_heading("2026 Presentation Deck Category Index & Structure", level=1)
    h_toc.runs[0].font.color.rgb = COLOR_PRIMARY
    h_toc.runs[0].font.size = Pt(15)
    
    toc_table = doc.add_table(rows=1, cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(toc_table)
    
    hdr_cells = toc_table.rows[0].cells
    hdr_titles = ["Category No.", "Presentation Module / Content Topic", "Slide Range"]
    for i, title_text in enumerate(hdr_titles):
        hdr_cells[i].text = title_text
        set_cell_bg(hdr_cells[i], COLOR_PRIMARY_HEX)
        set_cell_padding(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.size = Pt(9.5)

    categories_map = [
        ("Cat 1", "Executive Profile, Founder Vision & Brand Story", "Slides 01 – 05", list(range(1, 6))),
        ("Cat 2", "Chronicles of BNP Interiors (Year-by-Year Timeline 2006-2026)", "Slides 06 – 09", list(range(6, 10))),
        ("Cat 3", "Core Values, In-House ERP System & Growth Process", "Slides 10 – 11", list(range(10, 12))),
        ("Cat 4", "Infrastructure, HOMAG German Millwork Machinery & Turnkey Execution", "Slides 12 – 15", list(range(12, 16))),
        ("Cat 5", "Project Sectors & PAN India Geographic Footprint", "Slides 16 – 17", list(range(16, 18))),
        ("Cat 6", "Hospitality Division & Luxury Hotel Case Studies (Taj, Hyatt, Marriott, etc.)", "Slides 18 – 45", list(range(18, 46))),
        ("Cat 7", "High-End Members-Only Luxury Clubs (Jio World Drive BKC, Hyderabad)", "Slides 46 – 49", list(range(46, 50))),
        ("Cat 8", "Corporate Division & Commercial Headquarters (Palava, Ahmedabad, Pune, Kerala)", "Slides 50 – 74", list(range(50, 75))),
        ("Cat 9", "Malls, Retail, Educational & Healthcare Verticals (Dhirubhai Ambani School, Reliance)", "Slides 75 – 82", list(range(75, 83))),
        ("Cat 10", "Luxury Hi-End Residential Projects (Celebrity Bandra Mansion, Jamnagar RIL)", "Slides 83 – 87", list(range(83, 88))),
        ("Cat 11", "In Tune With Tomorrow & SOH Magazine Editorial Spotlight", "Slides 88 – 89", list(range(88, 90))),
        ("Cat 12", "Ongoing National Projects Showcase (Chennai, Delhi, Mumbai)", "Slides 90 – 91", list(range(90, 92))),
        ("Cat 13", "Corporate Head Office & Direct Contact Information", "Slide 92", [92])
    ]

    for cat_num, cat_title, s_range, _ in categories_map:
        row_cells = toc_table.add_row().cells
        row_cells[0].text = cat_num
        row_cells[1].text = cat_title
        row_cells[2].text = s_range
        
        set_cell_bg(row_cells[0], COLOR_BG_LIGHT_HEX)
        for i in range(3):
            set_cell_padding(row_cells[i], top=100, bottom=100, left=140, right=140)
            p = row_cells[i].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if i == 1:
                p.runs[0].font.bold = True

    widths = [Inches(1.0), Inches(4.5), Inches(1.3)]
    for row in toc_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_page_break()

    # --- CATEGORY BUILDER HELPER ---
    for cat_num, cat_title, s_range, slide_indices in categories_map:
        p_heading = doc.add_heading(f"{cat_num}: {cat_title} ({s_range})", level=1)
        p_heading.runs[0].font.color.rgb = COLOR_PRIMARY
        p_heading.runs[0].font.size = Pt(14)
        p_heading.paragraph_format.space_before = Pt(14)
        p_heading.paragraph_format.space_after = Pt(8)

        tbl = doc.add_table(rows=1, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl)
        
        hdr = tbl.rows[0].cells
        hdr[0].text = "Slide / Location"
        hdr[1].text = "Content Extracted from PDF Presentation"
        hdr[2].text = "Proposed Website Placement / Client Notes"
        
        for i in range(3):
            set_cell_bg(hdr[i], COLOR_PRIMARY_HEX)
            set_cell_padding(hdr[i], top=120, bottom=120, left=140, right=140)
            p = hdr[i].paragraphs[0]
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.runs[0].font.size = Pt(9)

        row_count = 0
        for s_idx in slide_indices:
            slide_obj = slides[s_idx - 1]
            slide_num = slide_obj['slide_num']
            lines = slide_obj['lines']
            
            if not lines:
                lines = ["[ Visual / Image-Based Slide — High Resolution Architecture & Interior Photography ]"]

            slide_label = f"Slide {slide_num:02d}"
            if len(lines) > 0 and len(lines[0]) < 50:
                slide_label += f"\n({lines[0]})"
                
            formatted_text = "\n".join(lines)

            row_cells = tbl.add_row().cells
            row_cells[0].text = sanitize_text(slide_label)
            row_cells[1].text = sanitize_text(formatted_text)
            row_cells[2].text = ""  # Client edit area

            bg_color = COLOR_ALT_ROW_HEX if row_count % 2 == 1 else "FFFFFF"
            set_cell_bg(row_cells[0], COLOR_BG_LIGHT_HEX)
            set_cell_bg(row_cells[1], bg_color)
            set_cell_bg(row_cells[2], "FFFFFF")

            for i in range(3):
                set_cell_padding(row_cells[i], top=100, bottom=100, left=140, right=140)
                p = row_cells[i].paragraphs[0]
                p.runs[0].font.size = Pt(9)
                if i == 0:
                    p.runs[0].font.bold = True
                    p.runs[0].font.color.rgb = COLOR_PRIMARY

            row_count += 1

        tbl_widths = [Inches(1.5), Inches(3.5), Inches(1.8)]
        for row in tbl.rows:
            for idx, w in enumerate(tbl_widths):
                row.cells[idx].width = w

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    output_filename = "BNP_Interiors_Profile_2026_Presentation_Content_Copy_Document.docx"
    doc.save(output_filename)
    print(f"Presentation Content Word Document generated successfully: {output_filename}")
    return output_filename

if __name__ == '__main__':
    build_presentation_word_document()
